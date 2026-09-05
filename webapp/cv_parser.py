"""
CV text extraction, self-contained for this service (deploy isolation -
see vacancy_source.py's docstring for why). Mirrors bot/bot.py's
handle_document logic exactly.
"""

import io

from docx import Document as DocxDocument
from pypdf import PdfReader


def parse_cv(filename: str, file_bytes: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif name.endswith(".docx"):
        docx_doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in docx_doc.paragraphs]
        for table in docx_doc.tables:
            for row in table.rows:
                paragraphs.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs)
    else:
        raise ValueError("Unsupported file type - only PDF and DOCX are accepted")
