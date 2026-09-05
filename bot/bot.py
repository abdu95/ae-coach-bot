import asyncio
import logging
import os
import io
import base64
import time

from dotenv import load_dotenv
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, WebAppInfo
from telegram.ext import (
    Application, CommandHandler, MessageHandler,
    CallbackQueryHandler, ContextTypes, filters,
)
from telegram.constants import ParseMode
from pypdf import PdfReader
from docx import Document as DocxDocument

load_dotenv()

import state
import coach
import formatter
import i18n
import vacancy_source

logging.basicConfig(
    format="%(asctime)s — %(name)s — %(levelname)s — %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

ADMIN_IDS = {
    int(x) for x in os.getenv("ADMIN_IDS", os.getenv("ADMIN_USER_ID", "0")).split(",") if x.strip()
}
FREE_LIMIT = int(os.getenv("FREE_LIMIT", "2"))
PILOT_CODE = os.getenv("PILOT_CODE", "school21")
# /start deep-link tags for tracking marketing campaigns -> users.source.
# e.g. t.me/<bot>?start=chashma for the chashma.uz half-marathon link.
MARKETING_SOURCES = {"chashma": "chashma_marathon"}
PILOT_QUOTA = int(os.getenv("PILOT_QUOTA", "10"))
PILOT_CAP = int(os.getenv("PILOT_CAP", "10"))
PAYME_ID = os.getenv("PAYME_ID", "")
PACKAGE_AMOUNT = int(os.getenv("PACKAGE_AMOUNT_TIYIN", "9900000"))  # 99,000 UZS in tiyin
PACKAGE_NAME = "10_checks"
MINI_APP_URL = os.getenv("MINI_APP_URL", "")


def effective_quota(user_id: int) -> int:
    override = state.get_quota_override(user_id)
    return override if override is not None else FREE_LIMIT


def can_run_check(user_id: int, user: dict) -> bool:
    return user["usage_count"] < effective_quota(user_id)


def action_button(label: str, callback: str) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([[InlineKeyboardButton(label, callback_data=callback)]])


MIN_LOADING_SECONDS = 1.5


async def ensure_min_display(started_at: float, minimum: float = MIN_LOADING_SECONDS) -> None:
    """Keeps a status message ('Reading your CV...', 'Analyzing...') on
    screen for at least `minimum` seconds even if the underlying call
    finishes faster, so it doesn't flash by unread."""
    elapsed = time.monotonic() - started_at
    if elapsed < minimum:
        await asyncio.sleep(minimum - elapsed)


LANG_BUTTONS = {
    "uz": ("🇺🇿 O'zbek", "lang_uz"),
    "ru": ("🇷🇺 Русский", "lang_ru"),
}


async def send_language_picker(message, hint_code: str | None) -> None:
    # language_code reflects the device's system language, not necessarily
    # the user's preferred chat language — used only to order the buttons.
    order = ["ru", "uz"] if hint_code == "ru" else ["uz", "ru"]
    buttons = [[InlineKeyboardButton(LANG_BUTTONS[c][0], callback_data=LANG_BUTTONS[c][1])] for c in order]
    await message.reply_text(
        "🌐 Choose your language / Tilni tanlang / Выберите язык",
        reply_markup=InlineKeyboardMarkup(buttons),
    )


def build_checkout_url(order_id: int, bot_username: str) -> str:
    raw = f"m={PAYME_ID};ac.order_id={order_id};a={PACKAGE_AMOUNT};c=https://t.me/{bot_username}"
    return "https://checkout.paycom.uz/" + base64.b64encode(raw.encode()).decode()


def checkout_url_for(user_id: int, bot_username: str) -> str:
    order_id = state.get_or_create_order(user_id, PACKAGE_AMOUNT, PACKAGE_NAME)
    return build_checkout_url(order_id, bot_username)


def out_of_checks_markup(lang: str, checkout_url: str) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton(i18n.pay_button(PACKAGE_AMOUNT, lang), url=checkout_url)],
        [InlineKeyboardButton(i18n.t("join_waitlist_button", lang), callback_data="join_waitlist")],
    ])


async def send_limit_reached(message, user_id: int, lang: str) -> None:
    state.log_event(user_id, "limit_reached")
    checkout_url = checkout_url_for(user_id, message.get_bot().username)
    await message.reply_text(
        i18n.limit_reached(effective_quota(user_id), lang),
        parse_mode=ParseMode.HTML,
        reply_markup=out_of_checks_markup(lang, checkout_url),
    )


# ── Commands ──────────────────────────────────────────────────────────────────

async def enroll_pilot(message, user_id: int, lang: str) -> None:
    if state.get_source(user_id) == "school21_pilot":
        await message.reply_text(i18n.t("pilot_already", lang))
        return
    if state.pilot_count() >= PILOT_CAP:
        await message.reply_text(i18n.t("pilot_full", lang))
        return
    state.set_quota_override(user_id, PILOT_QUOTA, source="school21_pilot")
    state.log_event(user_id, "pilot_enrolled")
    await message.reply_text(i18n.pilot_enrolled(PILOT_QUOTA, lang))


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_id = update.effective_user.id
    state.reset(user_id)
    state.log_event(user_id, "started")
    user = state.get(user_id)

    if context.args and context.args[0] == PILOT_CODE:
        await enroll_pilot(update.message, user_id, user["lang"])
    elif context.args and context.args[0] in MARKETING_SOURCES:
        state.set_source(user_id, MARKETING_SOURCES[context.args[0]])
        state.log_event(user_id, "marketing_source_start", metadata={"source": context.args[0]})

    if not user["lang"]:
        await send_language_picker(update.message, update.effective_user.language_code)
    elif not user["name"]:
        user["phase"] = "waiting_name"
        await update.message.reply_text(i18n.t("ask_name", user["lang"]))
    else:
        await update.message.reply_text(i18n.t("welcome", user["lang"]), parse_mode=ParseMode.HTML)


async def reset_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_id = update.effective_user.id
    had_cv = bool(state.get(user_id).get("cv_text"))
    state.reset(user_id)
    state.log_event(user_id, "reset", metadata={"had_cv": had_cv})
    await update.message.reply_text(i18n.t("reset_done", state.get(user_id)["lang"]))


async def language_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await send_language_picker(update.message, update.effective_user.language_code)


async def app_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Opens the Mini App (vacancy search -> apply -> track). No CV/JD
    needed in chat first - the app handles CV upload itself.

    Calling state.get() here (even though this handler doesn't use the
    session data) ensures a `users` row exists for this telegram_id.
    Without it, a user who only ever taps /app (never /start) would have
    no `users` row, and saving an application later would fail - the
    applications table has telegram_id REFERENCES users(telegram_id)."""
    user = state.get(update.effective_user.id)

    if not MINI_APP_URL:
        await update.message.reply_text(i18n.t("app_not_configured", user["lang"]))
        return
    await update.message.reply_text(
        i18n.t("app_intro", user["lang"]),
        reply_markup=InlineKeyboardMarkup([[
            InlineKeyboardButton(i18n.t("app_open_button", user["lang"]), web_app=WebAppInfo(url=MINI_APP_URL))
        ]]),
    )


async def jobs_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_id = update.effective_user.id
    user = state.get(user_id)
    state.log_event(user_id, "jobs_started")

    if not user.get("cv_text"):
        user["phase"] = "waiting_cv_for_jobs"
        await update.message.reply_text(i18n.t("jobs_need_cv", user["lang"]), parse_mode=ParseMode.HTML)
        return

    await start_job_title_suggestion(update.message, user_id)


async def stats(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if update.effective_user.id not in ADMIN_IDS:
        return
    acct = state.account_stats()
    pilot = state.pilot_stats()
    ev = state.event_stats(
        ["started", "name_provided", "cv_uploaded", "check_completed",
         "roadmap_requested", "limit_reached", "reset"]
    )
    reset_with_cv = state.reset_with_cv_count()
    chashma_count = state.source_count("chashma_marathon")

    def line(label: str, key: str) -> str:
        total, unique = ev[key]
        return f"{label} — Total {total} · Unique {unique}"

    await update.message.reply_text(
        f"📊 <b>Stats</b>\n\n"
        f"Unique users — {acct['unique_users']}\n"
        f"Activated (≥1 check) — {acct['activated_users']}\n"
        f"Total checks run — {acct['total_checks']}\n\n"
        f"{line('Started', 'started')}\n"
        f"{line('Name provided', 'name_provided')}\n"
        f"{line('CV uploaded', 'cv_uploaded')}\n"
        f"{line('Check completed', 'check_completed')}\n"
        f"{line('Roadmap requested', 'roadmap_requested')}\n"
        f"{line('Limit reached', 'limit_reached')}\n"
        f"{line('Reset', 'reset')} (wiped a stored CV — {reset_with_cv})\n\n"
        f"Waitlist — {acct['waitlist']}\n\n"
        f"🎓 School21 pilot — {pilot['users']} users · {pilot['checks']} checks · avg {pilot['avg']}/user\n"
        f"🏃 Chashma marathon — {chashma_count} users",
        parse_mode=ParseMode.HTML,
    )


async def grant_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if update.effective_user.id not in ADMIN_IDS:
        return
    if not context.args:
        await update.message.reply_text("Usage: /grant <telegram_id> [quota]")
        return
    try:
        target_id = int(context.args[0])
        quota = int(context.args[1]) if len(context.args) > 1 else PILOT_QUOTA
    except ValueError:
        await update.message.reply_text("Usage: /grant <telegram_id> [quota]")
        return
    if state.set_quota_override(target_id, quota):
        await update.message.reply_text(f"✅ Granted quota={quota} to {target_id}.")
    else:
        await update.message.reply_text(f"⚠️ {target_id} has never messaged the bot — nothing to grant.")


# ── CV upload ─────────────────────────────────────────────────────────────────

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_id = update.effective_user.id
    user = state.get(user_id)
    doc = update.message.document

    if not can_run_check(user_id, user):
        await send_limit_reached(update.message, user_id, user["lang"])
        return

    filename = doc.file_name.lower()
    if filename.endswith(".pdf"):
        file_kind = "pdf"
    elif filename.endswith(".docx"):
        file_kind = "docx"
    else:
        await update.message.reply_text(i18n.t("please_upload_cv", user["lang"]), parse_mode=ParseMode.HTML)
        return

    state.log_event(user_id, "cv_uploaded")
    msg = await update.message.reply_text(i18n.t("reading_cv", user["lang"]))
    started = time.monotonic()

    try:
        file = await context.bot.get_file(doc.file_id)
        cv_bytes = await file.download_as_bytearray()
        if file_kind == "pdf":
            reader = PdfReader(io.BytesIO(bytes(cv_bytes)))
            user["cv_text"] = "\n".join(page.extract_text() or "" for page in reader.pages)
        else:
            docx_doc = DocxDocument(io.BytesIO(bytes(cv_bytes)))
            paragraphs = [p.text for p in docx_doc.paragraphs]
            for table in docx_doc.tables:
                for row in table.rows:
                    paragraphs.append("\t".join(cell.text for cell in row.cells))
            user["cv_text"] = "\n".join(paragraphs)
        came_from_jobs = user["phase"] == "waiting_cv_for_jobs"
        user["phase"] = "waiting_jd" if not came_from_jobs else "idle"
    except Exception as e:
        logger.error(f"CV read error: {e}")
        await msg.edit_text(i18n.t("cv_read_error", user["lang"]))
        return

    await ensure_min_display(started)
    await msg.delete()

    if came_from_jobs:
        await start_job_title_suggestion(update.message, user_id)
        return

    await update.message.reply_text(
        i18n.t("cv_received", user["lang"]),
        parse_mode=ParseMode.HTML,
    )


# ── Vacancy suggestion flow (/jobs) ───────────────────────────────────────────

def work_setup_markup(lang: str) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton(i18n.t("work_setup_remote", lang), callback_data="worksetup_remote")],
        [InlineKeyboardButton(i18n.t("work_setup_hybrid", lang), callback_data="worksetup_hybrid")],
        [InlineKeyboardButton(i18n.t("work_setup_onsite", lang), callback_data="worksetup_onsite")],
    ])


async def start_job_title_suggestion(message, user_id: int) -> None:
    user = state.get(user_id)
    user["suggested_titles"] = []
    user["seen_companies"] = []
    user["search_count"] = 0
    user["current_vacancy"] = None
    user["chosen_vacancy"] = None
    user["job_title"] = ""
    user["location"] = ""
    user["work_setup"] = ""
    user["industry"] = ""

    loading = await message.reply_text(i18n.t("jobs_finding_titles", user["lang"]))
    started = time.monotonic()
    try:
        titles = await coach.suggest_job_titles(user["cv_text"])
    except Exception as e:
        logger.error(f"Job title suggestion error: {e}")
        await loading.edit_text(i18n.t("jobs_titles_failed", user["lang"]))
        return

    user["suggested_titles"] = titles
    user["phase"] = "waiting_job_title"
    await ensure_min_display(started)
    await loading.delete()

    buttons = [[InlineKeyboardButton(title, callback_data=f"jobtitle_{i}")] for i, title in enumerate(titles)]
    buttons.append([InlineKeyboardButton(i18n.t("jobs_regenerate_button", user["lang"]), callback_data="jobtitle_regenerate")])
    await message.reply_text(
        i18n.t("jobs_pick_title", user["lang"]),
        reply_markup=InlineKeyboardMarkup(buttons),
    )


async def run_vacancy_search(message, user_id: int) -> None:
    user = state.get(user_id)
    loading = await message.reply_text(i18n.t("jobs_searching", user["lang"]))
    started = time.monotonic()

    try:
        vacancies = await vacancy_source.search_vacancies(
            user["job_title"], user["location"] or "Any",
            user["work_setup"] or "Any", user["industry"] or "Any",
            seen_companies=user["seen_companies"],
        )
    except Exception as e:
        logger.error(f"Vacancy search error: {e}")
        await loading.edit_text(i18n.t("jobs_search_failed", user["lang"]))
        return

    user["search_count"] += 1

    if not vacancies:
        await loading.edit_text(i18n.t("jobs_no_match", user["lang"]))
        return

    vacancy = vacancies[0]
    try:
        score = await coach.score_vacancy(user["cv_text"], vacancy)
    except Exception as e:
        logger.error(f"Vacancy score error: {e}")
        await loading.edit_text(i18n.t("jobs_search_failed", user["lang"]))
        return

    user["current_vacancy"] = vacancy
    user["seen_companies"].append(vacancy["company"])
    await ensure_min_display(started)
    await loading.delete()

    card = formatter.vacancy_card(vacancy, score, user["search_count"], 3, user["lang"])
    buttons = [[InlineKeyboardButton(i18n.t("jobs_pick_button", user["lang"]), callback_data="vacancy_pick")]]
    if user["search_count"] < 3:
        buttons.append([InlineKeyboardButton(i18n.t("jobs_search_again_button", user["lang"]), callback_data="vacancy_again")])
    await message.reply_text(card, parse_mode=ParseMode.HTML, reply_markup=InlineKeyboardMarkup(buttons))


# ── JD text → run analysis ────────────────────────────────────────────────────

async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_id = update.effective_user.id
    user = state.get(user_id)
    text = update.message.text.strip()

    if user["phase"] == "waiting_name":
        name = text[:50]
        user["name"] = name
        user["phase"] = "idle"
        state.log_event(user_id, "name_provided")
        await update.message.reply_text(i18n.stats_and_privacy(name, user["lang"]))
        await update.message.reply_text(i18n.t("welcome", user["lang"]), parse_mode=ParseMode.HTML)
        return

    if user["phase"] == "waiting_location":
        user["location"] = text[:100]
        user["phase"] = "waiting_work_setup"
        await update.message.reply_text(
            i18n.t("jobs_ask_work_setup", user["lang"]),
            reply_markup=work_setup_markup(user["lang"]),
        )
        return

    if user["phase"] == "waiting_industry":
        user["industry"] = text[:100]
        await run_vacancy_search(update.message, user_id)
        return

    if user["phase"] != "waiting_jd":
        await update.message.reply_text(i18n.t("wrong_phase", user["lang"]))
        return

    if not can_run_check(user_id, user):
        await send_limit_reached(update.message, user_id, user["lang"])
        return

    if len(text) < 100:
        await update.message.reply_text(i18n.t("jd_too_short", user["lang"]))
        return

    user["jd"] = text
    user["phase"] = "analyzing"

    msg = await update.message.reply_text(
        i18n.t("analyzing", user["lang"]), parse_mode=ParseMode.HTML
    )
    started = time.monotonic()
    try:
        outputs = await coach.analyze_cv(user["jd"], user["cv_text"])
        user["outputs"] = outputs
        user["level"] = outputs["level"]["assessment"]
        user["phase"] = "step_1"
        user["usage_count"] += 1
        state.log_event(user_id, "check_completed")
    except Exception as e:
        logger.error(f"Analysis error: {e}")
        await msg.edit_text(i18n.t("analysis_failed", user["lang"]))
        return

    await ensure_min_display(started)
    await msg.delete()
    quota = effective_quota(user_id)
    remaining = max(0, quota - user["usage_count"])
    await update.message.reply_text(
        formatter.step_ats(outputs["ats"], user["lang"]) + f"\n\n{i18n.checks_left(remaining, quota, user['lang'])}",
        parse_mode=ParseMode.HTML,
        reply_markup=action_button(i18n.t("check_writing_button", user["lang"]), "step_2"),
    )


# ── Step callbacks ────────────────────────────────────────────────────────────

async def send_cv_fix(message, user_id: int, index: int) -> None:
    user = state.get(user_id)
    fixes = user["cv_fixes"]
    total = len(fixes)
    text = formatter.cv_fix_block(index, total, fixes[index - 1], user["lang"])

    if index < total:
        await message.reply_text(
            text,
            parse_mode=ParseMode.HTML,
            reply_markup=action_button(i18n.t("next_fix_button", user["lang"]), f"cvfix_{index + 1}"),
        )
    else:
        next_title = coach.roadmap_block_title(user["level"], 2)
        await message.reply_text(
            text,
            parse_mode=ParseMode.HTML,
            reply_markup=action_button(i18n.continue_button(next_title, user["lang"]), "step_5_2"),
        )


async def send_roadmap_item(message, user_id: int, item: int) -> None:
    user = state.get(user_id)
    level = user["level"]
    title = coach.roadmap_block_title(level, item)

    loading = await message.reply_text(
        i18n.roadmap_loading(item, title, user["lang"]),
        parse_mode=ParseMode.HTML,
    )
    started = time.monotonic()

    if title == "CV Fixes":
        try:
            fixes = await coach.generate_cv_fixes(level, user["jd"], user["cv_text"])
        except Exception as e:
            logger.error(f"Roadmap error (item {item}): {e}")
            await loading.edit_text(i18n.t("roadmap_failed", user["lang"]))
            return
        await ensure_min_display(started)
        await loading.delete()
        user["cv_fixes"] = fixes
        await send_cv_fix(message, user_id, index=1)
        return

    try:
        text = await coach.generate_roadmap(level, item, user["jd"], user["cv_text"])
    except Exception as e:
        logger.error(f"Roadmap error (item {item}): {e}")
        await loading.edit_text(i18n.t("roadmap_failed", user["lang"]))
        return

    await ensure_min_display(started)
    await loading.delete()
    formatted = formatter.step_roadmap_block(item, title, text, user["lang"])

    if item < coach.roadmap_max_item(level):
        next_title = coach.roadmap_block_title(level, item + 1)
        chunks = formatter.split_long(formatted)
        for chunk in chunks[:-1]:
            await message.reply_text(chunk, parse_mode=ParseMode.HTML)
        await message.reply_text(
            chunks[-1],
            parse_mode=ParseMode.HTML,
            reply_markup=action_button(i18n.continue_button(next_title, user["lang"]), f"step_5_{item + 1}"),
        )
    else:
        for chunk in formatter.split_long(formatted):
            await message.reply_text(chunk, parse_mode=ParseMode.HTML)
        quota = effective_quota(user_id)
        remaining = max(0, quota - user["usage_count"])
        done_text = i18n.analysis_done(remaining, quota, user["lang"])
        if remaining <= 0:
            checkout_url = checkout_url_for(user_id, message.get_bot().username)
            await message.reply_text(
                done_text, parse_mode=ParseMode.HTML,
                reply_markup=out_of_checks_markup(user["lang"], checkout_url),
            )
        else:
            await message.reply_text(done_text, parse_mode=ParseMode.HTML)


async def handle_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    query = update.callback_query
    await query.answer()
    user_id = update.effective_user.id
    user = state.get(user_id)
    action = query.data
    message = query.message

    await query.edit_message_reply_markup(reply_markup=None)

    if action in ("lang_uz", "lang_ru"):
        user["lang"] = "uz" if action == "lang_uz" else "ru"
        if user["name"]:
            await message.reply_text(i18n.t("welcome", user["lang"]), parse_mode=ParseMode.HTML)
        else:
            user["phase"] = "waiting_name"
            await message.reply_text(i18n.t("ask_name", user["lang"]))
        return

    if action == "join_waitlist":
        if state.join_waitlist(user_id, update.effective_user.username):
            state.log_event(user_id, "waitlist_joined")
        user["waitlisted"] = True
        await message.reply_text(i18n.t("waitlist_joined", user["lang"]))
        return

    if action == "jobtitle_regenerate":
        await start_job_title_suggestion(message, user_id)
        return

    if action.startswith("jobtitle_"):
        index = int(action.split("_")[1])
        user["job_title"] = user["suggested_titles"][index]
        user["phase"] = "waiting_location"
        await message.reply_text(i18n.t("jobs_ask_location", user["lang"]))
        return

    if action.startswith("worksetup_"):
        user["work_setup"] = action.split("_", 1)[1]
        user["phase"] = "waiting_industry"
        await message.reply_text(
            i18n.t("jobs_ask_industry", user["lang"]),
            reply_markup=action_button(i18n.t("jobs_industry_skip_button", user["lang"]), "industry_skip"),
        )
        return

    if action == "industry_skip":
        user["industry"] = ""
        await run_vacancy_search(message, user_id)
        return

    if action == "vacancy_again":
        if user["search_count"] >= 3:
            await message.reply_text(i18n.t("jobs_search_cap_reached", user["lang"]))
            return
        await run_vacancy_search(message, user_id)
        return

    if action == "vacancy_pick":
        user["chosen_vacancy"] = user["current_vacancy"]
        state.log_event(user_id, "vacancy_picked", metadata={"company": user["current_vacancy"]["company"]})
        await message.reply_text(i18n.t("jobs_picked", user["lang"]))
        return

    if not user.get("outputs"):
        await message.reply_text(i18n.t("session_expired", user["lang"]))
        return

    if action == "step_2":
        await message.reply_text(
            formatter.step_xyz(user["outputs"]["xyz"], user["lang"]),
            parse_mode=ParseMode.HTML,
            reply_markup=action_button(i18n.t("skill_gaps_button", user["lang"]), "step_3"),
        )

    elif action == "step_3":
        await message.reply_text(
            formatter.step_tools(user["outputs"]["tools"], user["lang"]),
            parse_mode=ParseMode.HTML,
            reply_markup=action_button(i18n.t("assess_level_button", user["lang"]), "step_4"),
        )

    elif action == "step_4":
        await message.reply_text(
            formatter.step_level(user["outputs"]["level"], user["lang"]),
            parse_mode=ParseMode.HTML,
            reply_markup=action_button(i18n.t("get_roadmap_button", user["lang"]), "step_5"),
        )

    elif action == "step_5":
        state.log_event(user_id, "roadmap_requested")
        await send_roadmap_item(message, user_id, item=1)

    elif action in ("step_5_2", "step_5_3", "step_5_4"):
        item = int(action.rsplit("_", 1)[1])
        await send_roadmap_item(message, user_id, item=item)

    elif action.startswith("cvfix_"):
        index = int(action.split("_")[1])
        await send_cv_fix(message, user_id, index=index)


async def error_handler(update: object, context: ContextTypes.DEFAULT_TYPE) -> None:
    logger.error("Unhandled exception", exc_info=context.error)


def main() -> None:
    token = os.getenv("TELEGRAM_TOKEN")
    if not token:
        raise ValueError("TELEGRAM_TOKEN not set")

    state.init_db()

    app = Application.builder().token(token).build()
    app.add_handler(CommandHandler("start", state.persisting(start)))
    app.add_handler(CommandHandler("reset", state.persisting(reset_cmd)))
    app.add_handler(CommandHandler("language", state.persisting(language_cmd)))
    app.add_handler(CommandHandler("app", state.persisting(app_cmd)))
    app.add_handler(CommandHandler("jobs", state.persisting(jobs_cmd)))
    app.add_handler(CommandHandler("stats", state.persisting(stats)))
    app.add_handler(CommandHandler("grant", state.persisting(grant_cmd)))
    app.add_handler(MessageHandler(filters.Document.ALL, state.persisting(handle_document)))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, state.persisting(handle_text)))
    app.add_handler(CallbackQueryHandler(state.persisting(handle_callback)))
    app.add_error_handler(error_handler)

    logger.info("Bot starting…")
    app.run_polling(drop_pending_updates=True)


if __name__ == "__main__":
    main()